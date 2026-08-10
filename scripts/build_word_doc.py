import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_callout(doc, title, text, bg_hex="F7F8FA", border_hex="C9A84C"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_hex)
    
    # Left border
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:left w:val="single" w:sz="36" w:space="0" w:color="{border_hex}"/>\n'
        f'  <w:top w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:bottom w:val="none"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(f"{title}\n")
    run_t.bold = True
    run_t.font.name = 'Inter'
    run_t.font.size = Pt(10)
    run_t.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)
    
    run_b = p.add_run(text)
    run_b.font.name = 'Inter'
    run_b.font.size = Pt(9.5)
    run_b.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_styled_heading(doc, text, level):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Inter'
    
    if level == 1:
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(8)
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)
        # Add thin line below
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="4" w:color="C9A84C"/></w:pBdr>')
        p._p.get_or_add_pPr().append(pBdr)
    elif level == 2:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x0F, 0x20, 0x38)
    elif level == 3:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run.font.size = Pt(11.5)
        run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    return p

def add_body_p(doc, text, bold_prefix=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.name = 'Inter'
        r_pre.font.size = Pt(10)
        r_pre.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    r = p.add_run(text)
    r.font.name = 'Inter'
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    return p

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.name = 'Inter'
        r_pre.font.size = Pt(9.5)
        r_pre.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    r = p.add_run(text)
    r.font.name = 'Inter'
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    return p

def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "0A1628")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(code_text)
    r.font.name = 'Consolas'
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_image_with_caption(doc, img_path, caption_text, width_in=6.0):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_in))
        
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(12)
        c_run = cp.add_run(caption_text)
        c_run.font.name = 'Inter'
        c_run.font.size = Pt(8.5)
        c_run.font.italic = True
        c_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

def format_table_headers_and_cells(tbl, col_widths, headers, data):
    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "0A1628")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.name = 'Inter'
            r.font.size = Pt(9)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
    # Data rows
    for r_idx, row_data in enumerate(data):
        row_cells = tbl.add_row().cells
        bg = "F7F8FA" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, cell_value in enumerate(row_data):
            row_cells[c_idx].text = str(cell_value)
            set_cell_background(row_cells[c_idx], bg)
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=120, right=120)
            p = row_cells[c_idx].paragraphs[0]
            for r in p.runs:
                r.font.name = 'Inter'
                r.font.size = Pt(8.5)
                r.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
                
    # Column widths
    for row in tbl.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = Inches(width)

def main():
    doc = Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    base_img_dir = "c:/Users/SHUBHAM/Desktop/acos/docs/images"
    logo_path = os.path.join(base_img_dir, "aravanta_logo.png")

    # ==================== COVER PAGE ====================
    p_cover_logo = doc.add_paragraph()
    p_cover_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cover_logo.paragraph_format.space_before = Pt(40)
    p_cover_logo.paragraph_format.space_after = Pt(20)
    if os.path.exists(logo_path):
        p_cover_logo.add_run().add_picture(logo_path, width=Inches(1.8))

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(4)
    r_t1 = p_title.add_run("Aravanta ")
    r_t1.bold = True
    r_t1.font.size = Pt(36)
    r_t1.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)
    r_t2 = p_title.add_run("CloudOS")
    r_t2.bold = True
    r_t2.font.size = Pt(36)
    r_t2.font.color.rgb = RGBColor(0xC9, 0xA8, 0x4C)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    r_sub = p_sub.add_run("CLOUD OPERATING PLATFORM")
    r_sub.font.size = Pt(13)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    p_desc = doc.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_desc.paragraph_format.space_after = Pt(40)
    r_desc = p_desc.add_run(
        "An open, self-service cloud platform for deploying, managing, and monitoring\n"
        "cloud-native applications and modern infrastructure from a unified command center."
    )
    r_desc.font.size = Pt(11)
    r_desc.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(0)
    r_meta = p_meta.add_run("Version 1.0 — Technical & Architectural Documentation\nJuly 2026 | Prepared by Aravanta Engineering Team")
    r_meta.font.size = Pt(9.5)
    r_meta.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    doc.add_page_break()

    # ==================== TABLE OF CONTENTS ====================
    add_styled_heading(doc, "Table of Contents", 1)
    
    toc_items = [
        ("Chapter 01 — Executive Summary", ["Vision & Mission", "Value Proposition"]),
        ("Chapter 02 — Platform Architecture", ["High-Level Architecture", "Architectural Layers", "Design Principles"]),
        ("Chapter 03 — Aravanta Service Catalog", ["ArvCompute", "ArvKube", "ArvStore", "ArvDB", "ArvRegistry", "ArvEdge", "ArvWatch", "ArvGate"]),
        ("Chapter 04 — Core Modules", ["Identity Module", "Dashboard Overview", "Infrastructure as Code", "CI/CD Pipeline", "Security Module", "Billing Engine"]),
        ("Chapter 05 — Technology Stack", ["Frontend Stack", "Backend Architecture", "DevOps & Infrastructure"]),
        ("Chapter 06 — API Reference", ["Authentication API", "Resource Management APIs"]),
        ("Chapter 07 — Deployment & Operations", ["Repository Structure", "Environment Promotion Pipeline", "Local Development"]),
        ("Chapter 08 — Monitoring & Observability", ["Metrics, Logs, and Traces", "Alerting Thresholds"]),
        ("Chapter 09 — Security Architecture", ["Defense-in-Depth", "Compliance & Governance"]),
        ("Chapter 10 — Roadmap & Portfolio Overview", ["Implementation Milestones", "Resume Project Summary"])
    ]

    for chap, subs in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(chap)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)
        for sub in subs:
            p_sub = doc.add_paragraph()
            p_sub.paragraph_format.left_indent = Inches(0.3)
            p_sub.paragraph_format.space_after = Pt(2)
            r_s = p_sub.add_run(f"• {sub}")
            r_s.font.size = Pt(9.5)
            r_s.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    doc.add_page_break()

    # ==================== CHAPTER 1 ====================
    add_styled_heading(doc, "Chapter 01 — Executive Summary", 1)
    
    add_callout(doc, "🌟 VISION STATEMENT", 
                "Aravanta CloudOS is a production-grade cloud operating platform designed to unify every aspect of modern cloud infrastructure into a single, intuitive experience. It empowers engineering teams to provision resources, deploy applications, manage Kubernetes workloads, automate infrastructure, and gain deep operational visibility — without switching between dozens of disconnected tools.",
                bg_hex="FEF9E7", border_hex="C9A84C")

    add_styled_heading(doc, "1.1 Vision & Mission", 2)
    add_body_p(doc, "The modern cloud ecosystem is fragmented. Engineers juggle separate consoles for compute, storage, databases, container orchestration, monitoring, and CI/CD. Each tool has its own learning curve, authentication mechanisms, and operational model. Aravanta CloudOS eliminates this fragmentation by acting as the operating system for cloud resources.")

    add_styled_heading(doc, "1.2 Value Proposition", 2)
    add_bullet(doc, "One-click application deployments with zero-downtime rollouts.", "For Developers: ")
    add_bullet(doc, "Built-in CI/CD pipelines connected directly to Git repositories.", "For Developers: ")
    add_bullet(doc, "Real-time log streaming and distributed tracing out of the box.", "For Developers: ")
    add_bullet(doc, "Centralized governance, audit logs, and granular RBAC across all services.", "For Organizations: ")
    add_bullet(doc, "Cost optimization through real-time usage analytics and budget alerts.", "For Organizations: ")

    # ==================== CHAPTER 2 ====================
    add_styled_heading(doc, "Chapter 02 — Platform Architecture", 1)
    add_body_p(doc, "Aravanta CloudOS follows a microservices-oriented, layered architecture. Each layer has clear responsibilities and communicates via well-defined RESTful and gRPC interfaces.")

    add_image_with_caption(doc, os.path.join(base_img_dir, "architecture_diagram.jpg"), 
                           "Figure 2.1 — Aravanta CloudOS High-Level Architecture Diagram", width_in=6.0)

    add_styled_heading(doc, "2.1 Architectural Layers", 2)
    t_arch = doc.add_table(rows=1, cols=3)
    format_table_headers_and_cells(t_arch, [1.5, 2.0, 3.0], 
                                  ["Layer", "Components", "Responsibility"],
                                  [
                                      ["Presentation Layer", "Web Dashboard, CLI, Mobile App", "User-facing interfaces providing graphical and command-line interactions."],
                                      ["API Gateway", "ArvGateway (Nginx + Middleware)", "Single entry point handling Auth, Rate Limiting, CORS, and TLS termination."],
                                      ["Application Services", "Auth, Compute, Kube, Storage, DB, CICD", "Domain-specific microservices implementing business logic and data persistence."],
                                      ["Orchestration Layer", "Container Runtime, Service Mesh", "Manages container lifecycles, service discovery, and inter-service mTLS."],
                                      ["Provider Abstraction", "Cloud Adapters Layer", "Translates Aravanta service requests to underlying cloud/bare-metal APIs."]
                                  ])

    # ==================== CHAPTER 3 ====================
    add_styled_heading(doc, "Chapter 03 — Aravanta Service Catalog", 1)
    add_body_p(doc, "Instead of copying cloud provider branding, Aravanta CloudOS introduces its own native, cloud-agnostic service abstractions:")

    services_data = [
        ("ArvCompute", "Compute Engine", "Virtual machines, auto-scaling groups, spot instances, instance templates, and SSH key management."),
        ("ArvKube", "Kubernetes Engine", "Fully managed Kubernetes clusters, node pools, rolling updates, pod inspection, and Helm charts."),
        ("ArvStore", "Object Storage", "High-durability object storage, bucket access policies, object versioning, and CDN integration."),
        ("ArvDB", "Managed Database", "Managed PostgreSQL & MySQL databases with automated HA failover, point-in-time recovery, and read replicas."),
        ("ArvRegistry", "Container Registry", "Private Docker image registry with automatic CVE scanning, image signing, and lifecycle rules."),
        ("ArvEdge", "Load Balancer & WAF", "L7 application load balancing, SSL/TLS termination, rate limiting, and edge DDoS protection."),
        ("ArvWatch", "Monitoring & Logs", "Prometheus metrics, Loki log aggregation, Jaeger tracing, Grafana dashboards, and Alertmanager alerts."),
        ("ArvGate", "Identity & Access", "JWT authentication, MFA (TOTP/FIDO2), RBAC policies, service accounts, and immutable audit logs.")
    ]

    t_serv = doc.add_table(rows=1, cols=3)
    format_table_headers_and_cells(t_serv, [1.5, 1.8, 3.2], ["Service Name", "Category", "Core Capabilities"], services_data)

    add_image_with_caption(doc, os.path.join(base_img_dir, "kubernetes_architecture.jpg"), 
                           "Figure 3.1 — ArvKube Managed Kubernetes Architecture", width_in=6.0)

    # ==================== CHAPTER 4 ====================
    add_styled_heading(doc, "Chapter 04 — Core Modules & Features", 1)
    
    add_styled_heading(doc, "4.1 Dashboard Command Center", 2)
    add_body_p(doc, "The unified dashboard provides real-time operational visibility into infrastructure health, active deployments, and resource utilization.")
    add_image_with_caption(doc, os.path.join(base_img_dir, "dashboard_mockup.jpg"), 
                           "Figure 4.1 — Aravanta CloudOS Unified Web Dashboard Mockup", width_in=6.0)

    add_styled_heading(doc, "4.2 Continuous Integration & Deployment (CI/CD)", 2)
    add_body_p(doc, "Automated deployment pipelines connect Git commits directly to staging and production Kubernetes environments.")
    add_image_with_caption(doc, os.path.join(base_img_dir, "cicd_pipeline.jpg"), 
                           "Figure 4.2 — Continuous Integration & Continuous Deployment (CI/CD) Workflow", width_in=6.0)

    add_styled_heading(doc, "4.3 Infrastructure as Code (IaC) Engine", 2)
    add_body_p(doc, "Native integration with Terraform allows users to define infrastructure declaratively, review plan outputs, apply changes safely, and detect configuration drift automatically.")

    # ==================== CHAPTER 5 ====================
    add_styled_heading(doc, "Chapter 05 — Technology Stack", 1)
    
    t_tech = doc.add_table(rows=1, cols=3)
    format_table_headers_and_cells(t_tech, [1.5, 2.0, 3.0], 
                                  ["Tier", "Technology", "Role & Rationale"],
                                  [
                                      ["Frontend", "React 18 + TypeScript + Tailwind CSS", "Modern, type-safe, responsive user interface with component reuse."],
                                      ["Backend API", "Python 3.11 + FastAPI", "High-performance async REST APIs with auto-generated OpenAPI documentation."],
                                      ["Database", "PostgreSQL 15 + Redis 7", "ACID transactional relational storage combined with fast in-memory caching."],
                                      ["Orchestration", "Docker + Kubernetes + Helm", "Standard container execution, cluster orchestration, and package management."],
                                      ["Infrastructure", "Terraform + AWS/Cloud SDKs", "Declarative IaC for multi-environment cloud resource management."],
                                      ["Monitoring", "Prometheus + Grafana + Loki", "Complete observability pillar stack for metrics, logs, and alerting."]
                                  ])

    # ==================== CHAPTER 6 ====================
    add_styled_heading(doc, "Chapter 06 — API Reference", 1)
    add_body_p(doc, "Aravanta CloudOS exposes RESTful endpoints secured via JWT bearer tokens.")
    
    add_styled_heading(doc, "6.1 Authentication Endpoints", 2)
    add_code_block(doc, "POST /api/v1/auth/login\nPOST /api/v1/auth/register\nPOST /api/v1/auth/mfa/verify")

    add_styled_heading(doc, "6.2 Compute & Cluster Provisioning Endpoints", 2)
    add_code_block(doc, "GET    /api/v1/compute/instances\nPOST   /api/v1/compute/instances\nDELETE /api/v1/compute/instances/{id}\n\nGET    /api/v1/kubernetes/clusters\nPOST   /api/v1/kubernetes/clusters\nGET    /api/v1/kubernetes/clusters/{id}/pods")

    # ==================== CHAPTER 7 ====================
    add_styled_heading(doc, "Chapter 07 — Deployment & Monorepo Structure", 1)
    add_code_block(doc, 
"""aravanta-cloudos/
├── frontend/             # React + TypeScript Web Application
├── backend/              # FastAPI Python Microservices
├── terraform/            # Infrastructure Provisioning Modules
├── kubernetes/           # K8s Manifests & Helm Charts
├── docker/               # Docker Compose for Local Development
├── monitoring/           # Prometheus, Grafana, Loki Configs
├── nginx/                # API Gateway Reverse Proxy Rules
├── scripts/              # Automated Setup & Seed Scripts
└── .github/workflows/    # CI/CD Pipeline Definitions""")

    # ==================== CHAPTER 8 ====================
    add_styled_heading(doc, "Chapter 08 — Monitoring & Observability", 1)
    add_body_p(doc, "Observability is built directly into every layer of Aravanta CloudOS using Prometheus metrics, Loki log streams, and Jaeger distributed traces.")
    add_image_with_caption(doc, os.path.join(base_img_dir, "monitoring_stack.jpg"), 
                           "Figure 8.1 — ArvWatch Observability & Alerting Stack", width_in=6.0)

    # ==================== CHAPTER 9 ====================
    add_styled_heading(doc, "Chapter 09 — Security Architecture", 1)
    add_body_p(doc, "The platform enforces a defense-in-depth model across network, application, and data layers.")
    add_image_with_caption(doc, os.path.join(base_img_dir, "security_architecture.jpg"), 
                           "Figure 9.1 — Defense-in-Depth Security Layers", width_in=6.0)

    # ==================== CHAPTER 10 ====================
    add_styled_heading(doc, "Chapter 10 — Roadmap & Resume Overview", 1)
    
    add_callout(doc, "💼 PORTFOLIO & RESUME SUMMARY", 
                "Aravanta CloudOS — A production-style cloud operating platform built with Python, React, Docker, Kubernetes, Terraform, and AWS/Cloud infrastructure. Provides secure authentication (JWT/MFA/RBAC), infrastructure provisioning (ArvCompute), managed Kubernetes orchestration (ArvKube), private container registry (ArvRegistry), object storage (ArvStore), database management (ArvDB), observability (ArvWatch), and CI/CD automation through a unified web interface.",
                bg_hex="EBF5FF", border_hex="3B82F6")

    output_doc_path = "c:/Users/SHUBHAM/Desktop/acos/docs/Aravanta_CloudOS_Documentation.docx"
    doc.save(output_doc_path)
    print(f"Successfully generated Word document at: {output_doc_path}")

if __name__ == "__main__":
    main()
